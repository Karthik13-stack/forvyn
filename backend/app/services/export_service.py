from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

class ExportService:

    def export_pdf(self, content: str) -> bytes:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os

        font_normal = 'Times-Roman'
        font_bold = 'Times-Bold'

        # Attempt to load a universal Indic font (Nirmala UI) on Windows or standard fallback for Linux/Docker
        nirmala_windows = 'C:/Windows/Fonts/Nirmala.ttc'
        nirmala_linux = '/usr/share/fonts/truetype/Nirmala.ttc'
        nirmala_path = nirmala_windows if os.name == 'nt' else nirmala_linux
        if os.path.exists(nirmala_path):
            try:
                pdfmetrics.registerFont(TTFont('Nirmala', nirmala_path, subfontIndex=0))
                font_normal = 'Nirmala'
                font_bold = 'Nirmala'
            except Exception as e:
                print(f"Failed to load Nirmala font: {e}")

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='LegalNormal', parent=styles['Normal'], fontName=font_normal, fontSize=12, leading=14, alignment=TA_JUSTIFY, spaceAfter=12))
        styles.add(ParagraphStyle(name='LegalTitle', parent=styles['Heading1'], fontName=font_bold, fontSize=16, alignment=TA_CENTER, spaceAfter=24, leading=20))
        styles.add(ParagraphStyle(name='LegalHeading', parent=styles['Heading2'], fontName=font_bold, fontSize=12, alignment=TA_LEFT, spaceAfter=12, leading=14))
        styles.add(ParagraphStyle(name='LegalClause', parent=styles['Normal'], fontName=font_normal, fontSize=12, leading=14, alignment=TA_JUSTIFY, leftIndent=36, firstLineIndent=-36, spaceAfter=12))
        styles.add(ParagraphStyle(name='LegalSubClause', parent=styles['Normal'], fontName=font_normal, fontSize=12, leading=14, alignment=TA_JUSTIFY, leftIndent=72, firstLineIndent=-36, spaceAfter=12))
        story = []
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if para.startswith('# '):
                text = para[2:].strip()
                story.append(Paragraph(text, styles['LegalTitle']))
            elif para.startswith('## '):
                text = para[3:].strip()
                story.append(Paragraph(text, styles['LegalHeading']))
            elif para.startswith('**'):
                match = re.match('\\*\\*([\\d.]+)\\*\\*\\s*(.*)', para, re.DOTALL)
                if match:
                    clause_id = match.group(1)
                    clause_text = match.group(2)
                    dots = clause_id.count('.')
                    formatted_text = self._format_markdown(clause_text)
                    full_line = f'<b>{clause_id}</b>\t{formatted_text}'
                    full_line = f'<b>{clause_id}</b> {formatted_text}'
                    if dots >= 2:
                        story.append(Paragraph(full_line, styles['LegalSubClause']))
                    else:
                        story.append(Paragraph(full_line, styles['LegalClause']))
                else:
                    formatted = self._format_markdown(para)
                    story.append(Paragraph(formatted, styles['LegalNormal']))
            else:
                formatted = self._format_markdown(para)
                story.append(Paragraph(formatted, styles['LegalNormal']))
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def export_docx(self, content: str) -> bytes:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(12)
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            p = None
            if para.startswith('# '):
                p = doc.add_paragraph(para[2:].strip())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(16)
                p.paragraph_format.space_after = Pt(24)
            elif para.startswith('## '):
                p = doc.add_paragraph(para[3:].strip())
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(12)
                p.paragraph_format.keep_with_next = True
            elif para.startswith('**'):
                match = re.match('\\*\\*([\\d.]+)\\*\\*\\s*(.*)', para, re.DOTALL)
                if match:
                    clause_id = match.group(1)
                    body_text = match.group(2)
                    clean_body = body_text.replace('**', '')
                    p = doc.add_paragraph()
                    run_id = p.add_run(f'{clause_id}\t')
                    run_id.bold = True
                    p.add_run(clean_body)
                    dots = clause_id.count('.')
                    pf = p.paragraph_format
                    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    indent_step = 0.5
                    hanging = 0.5
                    level = max(0, dots - 1)
                    left_indent = level * indent_step + hanging
                    pf.left_indent = Inches(left_indent)
                    pf.first_line_indent = Inches(-hanging)
                    pf.tab_stops.add_tab_stop(Inches(left_indent))
                else:
                    clean = para.replace('**', '')
                    p = doc.add_paragraph(clean)
            else:
                clean = para.replace('**', '')
                doc.add_paragraph(clean)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _format_markdown(self, text: str) -> str:
        text = re.sub('\\*\\*(.*?)\\*\\*', '<b>\\1</b>', text)
        text = re.sub('\\*(.*?)\\*', '<i>\\1</i>', text)
        text = text.replace('\n', '<br/>')
        return text